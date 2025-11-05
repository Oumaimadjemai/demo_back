
from docx import Document
from .models import *
from .serializers import *
from django.db.models import Sum,Count,Value, DecimalField
from rest_framework import generics, filters, status
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.response import Response
from .filters import *
from django.db.models.functions import Coalesce

from django.db.models import Sum, Count, Value, DecimalField, IntegerField, Q
from django.db.models.functions import Coalesce
from rest_framework.response import Response
from rest_framework import generics, filters
from django_filters.rest_framework import DjangoFilterBackend

from .models import VenteFacilite
from .serializers import VenteFaciliteSerializer
from .filters import VenteFaciliteFilter  # if you already have it
from rest_framework.decorators import action



from django.db.models import Count, Sum, F, DecimalField, IntegerField, Value, Q
from django.db.models.functions import Coalesce
from decimal import Decimal
from rest_framework import generics, filters
from rest_framework.response import Response
from django_filters.rest_framework import DjangoFilterBackend
from .models import VenteFacilite
from .serializers import VenteFaciliteSerializer
from .filters import VenteFaciliteFilter

class VenteFaciliteCreateAPIView(generics.ListCreateAPIView):
    queryset = VenteFacilite.objects.all().select_related('client', 'utilisateur').order_by('-id')
    serializer_class = VenteFaciliteSerializer
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    filterset_class = VenteFaciliteFilter

    # ✅ Search client by any name or CCP
    search_fields = [
        'client__nom_famille_fr',
        'client__prenom_fr',
        'client__nom_famille_ar',
        'client__prenom_ar',
        'client__ccp',
        'client__cle'
    ]

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        serializer = self.get_serializer(queryset, many=True)

        # ✅ Base stats
        stats = queryset.aggregate(
            total_ventes=Count('id'),
            total_montant_total=Coalesce(
                Sum('montant_total', output_field=DecimalField()),
                Value(0, output_field=DecimalField())
            ),
            total_montant_verse=Coalesce(
                Sum('montant_verse', output_field=DecimalField()),
                Value(0, output_field=DecimalField())
            ),
            total_montant_restant=Coalesce(
                Sum('montant_restant', output_field=DecimalField()),
                Value(0, output_field=DecimalField())
            ),
        )

        # ✅ Calcul du coût d'achat total
        cout_achat_total = queryset.aggregate(
            total_achat=Coalesce(
                Sum(
                    F('lignes__quantite') * F('lignes__produit__prix_achat'),
                    output_field=DecimalField()
                ),
                Value(0, output_field=DecimalField())
            )
        )['total_achat'] or Decimal('0')

        # ✅ Bénéfice total = Montant ventes - coût d'achat
        total_montant_total = stats['total_montant_total'] or Decimal('0')
        benefice_total = total_montant_total - cout_achat_total

        # ✅ Pourcentage bénéfice = (bénéfice / coût d'achat) * 100
        pourcentage_benefice = (
            (benefice_total / cout_achat_total * 100)
            if cout_achat_total > 0
            else 0
        )

        # ✅ Extra stats pour produit sélectionné
        produit_id = request.query_params.get('produit')
        quantite_vendue = 0
        montant_total_produit = 0
        benefice_produit = 0
        pourcentage_benefice_produit = 0

        if produit_id:
            produit_stats = queryset.aggregate(
                quantite_vendue=Coalesce(
                    Sum(
                        'lignes__quantite',
                        filter=Q(lignes__produit_id=produit_id),
                        output_field=IntegerField()
                    ),
                    Value(0, output_field=IntegerField())
                ),
                montant_total_produit=Coalesce(
                    Sum(
                        'lignes__sous_total',
                        filter=Q(lignes__produit_id=produit_id),
                        output_field=DecimalField()
                    ),
                    Value(0, output_field=DecimalField())
                ),
                cout_achat_produit=Coalesce(
                    Sum(
                        F('lignes__quantite') * F('lignes__produit__prix_achat'),
                        filter=Q(lignes__produit_id=produit_id),
                        output_field=DecimalField()
                    ),
                    Value(0, output_field=DecimalField())
                )
            )

            quantite_vendue = produit_stats['quantite_vendue'] or 0
            montant_total_produit = produit_stats['montant_total_produit'] or 0
            cout_achat_produit = produit_stats['cout_achat_produit'] or Decimal('0')

            benefice_produit = montant_total_produit - cout_achat_produit
            pourcentage_benefice_produit = (
                (benefice_produit / cout_achat_produit * 100)
                if cout_achat_produit > 0 else 0
            )

        return Response({
            "ventes": serializer.data,
            **stats,
            "quantite_vendue_produit": quantite_vendue,
            "montant_total_produit": montant_total_produit,
            "cout_achat_total": cout_achat_total,
            "benefice_total": benefice_total,
            "pourcentage_benefice_total": round(pourcentage_benefice, 2),
            "benefice_produit": benefice_produit,
            "pourcentage_benefice_produit": round(pourcentage_benefice_produit, 2),
        })

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from decimal import Decimal
from django.utils.timezone import now
from django.db.models import Sum


from decimal import Decimal
from django.db import transaction
from django.db.models import Sum
from django.utils.timezone import now
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status

from .models import VenteFacilite, PaiementVente
from .serializers import PaiementVenteSerializer


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.db import transaction
from django.db.models import Sum
from decimal import Decimal
from django.utils.timezone import now
from .models import VenteFacilite, PaiementVente


class VentePaiementAPIView(APIView):
    """
    GET    -> Historique des ventes avec filtres (client, year, month_index, status)
    POST   -> Effectuer un paiement pour une vente
    PATCH  -> Changer le statut d’un paiement existant
    """

    def get(self, request, pk=None):
        client_id = request.query_params.get("client")
        year_filter = request.query_params.get("year")  # ex: 2025
        month_index_filter = request.query_params.get("month")  # 0-based
        status_filter = request.query_params.get("status")  # paid, unpaid, pending

        ventes = VenteFacilite.objects.all()

        if pk:
            ventes = ventes.filter(id=pk)
        elif client_id:
            ventes = ventes.filter(client_id=client_id)

        if not ventes.exists():
            return Response({"error": "Aucune vente trouvée."}, status=status.HTTP_404_NOT_FOUND)

        history = []

        for vente in ventes:
            paiements_qs = vente.paiements.all().order_by("date_paiement")

            # --- Apply filters on paiements ---
            if status_filter:
                paiements_qs = paiements_qs.filter(status=status_filter)
            if year_filter:
                paiements_qs = paiements_qs.filter(date_paiement__year=int(year_filter))
            if month_index_filter:
                paiements_qs = paiements_qs.filter(month_index=int(month_index_filter))

            if not paiements_qs.exists() and (status_filter or month_index_filter or year_filter):
                # Skip ventes that don't match filter
                continue

            mois_status = vente.mois_status or ["pending"] * vente.nombre_mois
            mois_status_filtered = mois_status[:]

            # --- If filtering by status, hide other months ---
            if status_filter:
                mois_status_filtered = [
                    s if s == status_filter else None for s in mois_status
                ]

            # --- If filtering by month index, hide other months ---
            if month_index_filter:
                month_index_filter = int(month_index_filter)
                mois_status_filtered = [
                    s if i == month_index_filter else None
                    for i, s in enumerate(mois_status_filtered)
                ]

            unpaid_count = vente.paiements.filter(status="pending").count()

            history.append({
                "vente_id": vente.id,
                "client_detail": {
                    "id": vente.client.id,
                    "nom_famille_ar": vente.client.nom_famille_ar,
                    "prenom_ar": vente.client.prenom_ar,
                    "nom_famille_fr": vente.client.nom_famille_fr,
                    "prenom_fr": vente.client.prenom_fr,
                    "ccp": vente.client.ccp,
                    "cle": vente.client.cle,
                    "telephone": vente.client.telephone,
                },
                "montant_mensuel": float(vente.montant_mensuel),
                "montant_par_mois": vente.montants_par_mois,
                "nombre_mois": vente.nombre_mois,
                "date_debut": vente.date_debut.strftime("%Y-%m-%d"),
                "mois_status": mois_status_filtered,
                "total_verse_effectif": float(vente.montant_paye_effectif),
                "reste_vente": float(vente.montant_restant),
                "reste_global": float(
                    VenteFacilite.objects.filter(client=vente.client)
                    .aggregate(total_restant=Sum('montant_restant'))['total_restant'] or 0
                ),
                "solde_regle": vente.montant_restant <= 0,
                "unpaid_count": unpaid_count,
                "paiements": [
                    {
                        "id": p.id,
                        "montant": float(p.montant),
                        "date_paiement": p.date_paiement.isoformat(),
                        "month_index": p.month_index,
                        "status": p.status
                    }
                    for p in paiements_qs
                ]
            })

        return Response(history, status=status.HTTP_200_OK)


    def patch(self, request, pk=None):
        """
        PATCH /vente/vente/payer/<paiement_id>/
        Body: {"status": "pending" | "paid" | "unpaid"}
        """
        paiement_id = pk
        new_status = request.data.get("status")

        if new_status not in ["pending", "paid", "unpaid"]:
            return Response(
                {"error": "Status invalide. Choisissez 'pending', 'paid' ou 'unpaid'."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            paiement = PaiementVente.objects.select_related("vente").get(id=paiement_id)
        except PaiementVente.DoesNotExist:
            return Response({"error": "Paiement introuvable."}, status=status.HTTP_404_NOT_FOUND)

        old_status = paiement.status
        if old_status == new_status:
            return Response(
                {"message": "Aucune modification : le statut est identique."},
                status=status.HTTP_200_OK,
            )

        vente = paiement.vente

        with transaction.atomic():
            montant_change = paiement.montant

            # --- Mettre à jour montants ---
            if old_status == "paid" and new_status in ["pending", "unpaid"]:
                vente.montant_paye_effectif -= montant_change
                vente.montant_restant += montant_change
            elif old_status in ["pending", "unpaid"] and new_status == "paid":
                vente.montant_paye_effectif += montant_change
                vente.montant_restant -= montant_change

            # --- Mettre à jour mois_status ---
            if paiement.month_index is not None:
                mois_status = vente.mois_status or ["pending"] * vente.nombre_mois
                if 0 <= paiement.month_index < len(mois_status):
                    mois_status[paiement.month_index] = new_status
                    vente.mois_status = mois_status

            # --- Sauvegarder paiement ---
            paiement.status = new_status
            paiement.save(update_fields=["status"])

            # --- Vérifier si la vente est réglée ---
            vente.solde_regle = vente.montant_restant <= 0
            vente.save(update_fields=[
                "montant_paye_effectif",
                "montant_restant",
                "mois_status",
                "solde_regle"
            ])

        reste_global = (
            VenteFacilite.objects.filter(client=vente.client)
            .aggregate(total_restant=Sum("montant_restant"))["total_restant"] or 0
        )
        unpaid_count = vente.paiements.filter(status="pending").count()

        return Response({
            "message": f"Paiement #{paiement.id} passé de {old_status} à {new_status}.",
            "vente_id": vente.id,
            "total_verse_effectif": float(vente.montant_paye_effectif),
            "reste_vente": float(vente.montant_restant),
            "reste_global": float(reste_global),
            "solde_regle": vente.solde_regle,
            "unpaid_count": unpaid_count,
        }, status=status.HTTP_200_OK)

    def post(self, request, pk):
        """
        POST /vente/vente/payer/<vente_id>/
        Body: {"montant": 5000, "month": 0-11, "status": "pending|paid|unpaid"}
        """
        try:
            vente = VenteFacilite.objects.get(pk=pk)
        except VenteFacilite.DoesNotExist:
            return Response({"error": "Vente introuvable"}, status=status.HTTP_404_NOT_FOUND)

        montant = Decimal(request.data.get("montant", vente.montant_mensuel))
        calendar_month_index = int(request.data.get("month", 0))  # 0-based
        status_str = request.data.get("status", "pending")

        if vente.montant_restant <= 0:
            return Response({"message": "La dette est déjà réglée."}, status=status.HTTP_400_BAD_REQUEST)

        paiement = min(montant, vente.montant_restant)

        # ✅ Update vente
        vente.montant_restant -= paiement
        vente.montant_paye_effectif += paiement
        vente.last_payment_date = now().date()

        mois_status = vente.mois_status or ["pending"] * vente.nombre_mois
        if 0 <= calendar_month_index < len(mois_status):
            mois_status[calendar_month_index] = status_str

        vente.mois_status = mois_status
        vente.solde_regle = vente.montant_restant <= 0
        vente.save()

        # ✅ Save paiement
        PaiementVente.objects.create(
            vente=vente,
            montant=paiement,
            month_index=calendar_month_index,
            status=status_str
        )

        return self.get(request, pk)
    
from collections import defaultdict
from decimal import Decimal
from django.utils.timezone import now
from rest_framework.response import Response
from rest_framework import status
from rest_framework.views import APIView
from django.core.files.storage import default_storage
from ventes.models import Client, VenteFacilite, PaiementVente


from collections import defaultdict
from decimal import Decimal
from django.utils.timezone import now
from django.db import transaction
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.core.files.storage import default_storage

from ventes.models import Client, VenteFacilite, PaiementVente


# class ImportPaiementsTxtAPIView(APIView):
#     """
#     POST /vente/import-paiements-txt/
#     Body: { "file": <fichier.txt> }
#     """

#     def handle_file(self, file_path):
#         paiements_effectues = []
#         erreurs = []

#         # 👇 Grouping (ccp + month) -> total amount
#         grouped_paiements = defaultdict(Decimal)

#         with open(file_path, "r", encoding="utf-8") as f:
#             for line in f:
#                 try:
#                     ccp = line[0:10].strip()
#                     name = line[10:35].strip()
#                     montant_str = line[35:49].strip()
#                     date_str = line[49:59].strip()  # format DD/MM/YYYY

#                     # Nettoyage du nom
#                     if name.startswith("M."):
#                         name = name[2:].strip()
#                     elif name.lower().startswith("melle."):
#                         name = name[6:].strip()

#                     # Conversion montant
#                     try:
#                         montant = Decimal(montant_str)
#                     except Exception:
#                         erreurs.append({
#                             "ligne": line.strip(),
#                             "error": f"Montant invalide: {montant_str}"
#                         })
#                         continue

#                     # Extraire le mois depuis la date
#                     try:
#                         month_index = int(date_str[3:5]) - 1  # "MM" → index (0-based)
#                     except Exception:
#                         erreurs.append({
#                             "ligne": line.strip(),
#                             "error": f"Date invalide: {date_str}"
#                         })
#                         continue

#                     # 👇 Group par client+mois
#                     grouped_paiements[(ccp, month_index)] += montant

#                 except Exception as e:
#                     erreurs.append({
#                         "ligne": line.strip(),
#                         "error": str(e)
#                     })

#         # 👇 Appliquer les paiements groupés
#         for (ccp, month_index), total_montant in grouped_paiements.items():
#             try:
#                 # 1️⃣ Client
#                 try:
#                     client = Client.objects.get(ccp=ccp)
#                 except Client.DoesNotExist:
#                     erreurs.append({
#                         "ccp": ccp,
#                         "error": f"Client introuvable CCP={ccp}"
#                     })
#                     continue

#                 # 2️⃣ Vente active
#                 vente = VenteFacilite.objects.filter(client=client, solde_regle=False).first()
#                 if not vente:
#                     erreurs.append({
#                         "client": client.id,
#                         "error": "Pas de vente active trouvée"
#                     })
#                     continue

#                 # 3️⃣ Paiement
#                 with transaction.atomic():
#                     PaiementVente.objects.create(
#                         vente=vente,
#                         montant=total_montant,
#                         month_index=month_index,
#                         status="paid",
#                         date_paiement=now()
#                     )

#                     # ✅ Mise à jour de la vente (sans toucher mois_status)
#                     vente.montant_paye_effectif += total_montant
#                     vente.montant_restant = max(Decimal(0), vente.montant_restant - total_montant)
#                     if vente.montant_restant == 0:
#                         vente.solde_regle = True
#                     vente.last_payment_date = now().date()
#                     vente.save()

#                 paiements_effectues.append({
#                     "client": str(client),
#                     "vente": vente.id,
#                     "paiement": str(total_montant),
#                     "mois": month_index + 1
#                 })

#             except Exception as e:
#                 erreurs.append({
#                     "ccp": ccp,
#                     "error": str(e)
#                 })

#         return {"paiements_effectues": paiements_effectues, "erreurs": erreurs}

#     def post(self, request, *args, **kwargs):
#         file = request.FILES.get("file")
#         if not file:
#             return Response({"error": "Aucun fichier fourni"}, status=status.HTTP_400_BAD_REQUEST)

#         file_path = default_storage.save(file.name, file)
#         try:
#             result = self.handle_file(file_path)
#         finally:
#             default_storage.delete(file_path)

#         return Response(result, status=status.HTTP_201_CREATED)
import io
from collections import defaultdict
from decimal import Decimal
from django.db import transaction
from django.utils.timezone import now
from django.core.files.storage import default_storage
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from ventes.models import Client, VenteFacilite, PaiementVente  # adjust import as needed

class ImportPaiementsTxtAPIView(APIView):
    """
    POST /vente/import-paiements-txt/
    Body: { "file": <fichier.txt> }
    """

    def handle_file(self, file_path):
        if not default_storage.exists(file_path):
            return {"paiements_effectues": [], "erreurs": [{"error": "Fichier introuvable"}]}

        paiements_effectues = []
        erreurs = []

        # Grouping (ccp + month) -> total amount
        grouped_paiements = defaultdict(Decimal)

        # Open binary file and wrap as text
        with default_storage.open(file_path, "rb") as f:
            text_file = io.TextIOWrapper(f, encoding="utf-8")
            for line in text_file:
                try:
                    ccp = line[2:10].strip()
                    name = line[10:35].strip()
                    montant_str = line[35:49].strip()
                    date_str = line[53:63].strip()  # format DD/MM/YYYY

                    # Nettoyage du nom
                    if name.startswith("M."):
                        name = name[2:].strip()
                    elif name.lower().startswith("melle."):
                        name = name[6:].strip()

                    # Conversion montant
                    try:
                        montant = Decimal(montant_str)
                    except Exception:
                        erreurs.append({
                            "ligne": line.strip(),
                            "error": f"Montant invalide: {montant_str}"
                        })
                        continue

                    # Extraire le mois depuis la date
                    try:
                        month_index = int(date_str[3:5]) - 1  # "MM" → index (0-based)
                    except Exception:
                        erreurs.append({
                            "ligne": line.strip(),
                            "error": f"Date invalide: {date_str}"
                        })
                        continue

                    # Group par client+mois
                    grouped_paiements[(ccp, month_index)] += montant

                except Exception as e:
                    erreurs.append({
                        "ligne": line.strip(),
                        "error": str(e)
                    })

        # Appliquer les paiements groupés
        for (ccp, month_index), total_montant in grouped_paiements.items():
            try:
                # 1️⃣ Client
                try:
                    client = Client.objects.get(ccp=ccp)
                except Client.DoesNotExist:
                    erreurs.append({
                        "ccp": ccp,
                        "error": f"Client introuvable CCP={ccp}"
                    })
                    continue

                # 2️⃣ Vente active
                vente = VenteFacilite.objects.filter(client=client, solde_regle=False).first()
                if not vente:
                    erreurs.append({
                        "client": client.id,
                        "error": "Pas de vente active trouvée"
                    })
                    continue

                # 3️⃣ Paiement
                with transaction.atomic():
                    PaiementVente.objects.create(
                        vente=vente,
                        montant=total_montant,
                        month_index=month_index,
                        status="paid",
                        date_paiement=now()
                    )

                    # Mise à jour de la vente (sans toucher mois_status)
                    vente.montant_paye_effectif += total_montant
                    vente.montant_restant = max(Decimal(0), vente.montant_restant - total_montant)
                    if vente.montant_restant == 0:
                        vente.solde_regle = True
                    vente.last_payment_date = now().date()
                    vente.save()

                paiements_effectues.append({
                    "client": str(client),
                    "vente": vente.id,
                    "paiement": str(total_montant),
                    "mois": month_index + 1
                })

            except Exception as e:
                erreurs.append({
                    "ccp": ccp,
                    "error": str(e)
                })

        return {"paiements_effectues": paiements_effectues, "erreurs": erreurs}

    def post(self, request, *args, **kwargs):
        file = request.FILES.get("file")
        if not file:
            return Response({"error": "Aucun fichier fourni"}, status=status.HTTP_400_BAD_REQUEST)

        file_path = default_storage.save(file.name, file)
        try:
            result = self.handle_file(file_path)
        finally:
            if default_storage.exists(file_path):
                default_storage.delete(file_path)

        return Response(result, status=status.HTTP_201_CREATED)

from django.db import transaction
class VenteFaciliteRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = VenteFacilite.objects.all()
    serializer_class = VenteFaciliteSerializer

    @transaction.atomic
    def destroy(self, request, *args, **kwargs):
        """
        Suppression d'une vente avec rollback automatique en cas d'erreur
        et restitution du stock produit.
        """
        instance = self.get_object()

        try:
            # 🔄 Restaurer les produits avant suppression
            for ligne in instance.lignes.all():
                produit = ligne.produit

                # Restituer la quantité
                produit.quantite += ligne.quantite

                # Restituer les codes-barres
                if ligne.codes_barres_utilises:
                    produit.codes_barres = ligne.codes_barres_utilises + produit.codes_barres

                produit.save()

            # Ensuite, supprimer la vente (lignes supprimées automatiquement via CASCADE)
            self.perform_destroy(instance)

            return Response(
                {"success": "Vente supprimée avec succès et stock restauré"},
                status=status.HTTP_204_NO_CONTENT
            )

        except Exception as e:
            transaction.set_rollback(True)
            return Response(
                {"error": f"Suppression annulée : {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def payer_mensualite(self, request, pk=None):
        vente = self.get_object()
        montant_paye = vente.effectuer_paiement_mensuel()
        if montant_paye > 0:
            return Response({
                "message": f"Paiement de {montant_paye} effectué",
                "montant_restant": vente.montant_restant,
                "montant_verse": vente.montant_verse,
                "last_payment_date": vente.last_payment_date,
                "dette_totale_client": vente.dette_totale_client,
                "dette_actuelle_client": vente.dette_actuelle_client
            })
        return Response({"message": "Vente déjà totalement payée"}, status=status.HTTP_400_BAD_REQUEST)

class VenteCacheCreateAPIView(generics.CreateAPIView):
    queryset = VenteCache.objects.all()
    serializer_class = VenteCacheSerializer



class VenteCacheListAPIView(generics.ListAPIView):
    queryset = VenteCache.objects.all()
    serializer_class = VenteCacheSerializer
    filter_backends = [DjangoFilterBackend, filters.SearchFilter]
    filterset_class = VenteCacheFilter

    # search client by any name
    search_fields = [
    'client__nom_famille_fr',
    'client__prenom_fr',
    'client__nom_famille_ar',
    'client__prenom_ar',
    'client__ccp',
    'client__cle'
]


    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        serializer = self.get_serializer(queryset, many=True)

        total_montant = queryset.aggregate(Sum("montant_total"))["montant_total__sum"] or 0
        total_ventes = queryset.count()
        produit_id = request.query_params.get('produit')
        quantite_vendue = 0
        montant_total_produit = 0

        if produit_id:
            produit_stats = queryset.aggregate(
                quantite_vendue=Coalesce(
                    Sum(
                        'lignes__quantite',
                        filter=Q(lignes__produit_id=produit_id),
                        output_field=IntegerField()
                    ),
                    Value(0, output_field=IntegerField())
                ),
                montant_total_produit=Coalesce(
                    Sum(
                        'lignes__sous_total',
                        filter=Q(lignes__produit_id=produit_id),
                        output_field=DecimalField()
                    ),
                    Value(0, output_field=DecimalField())
                ),
            )

            quantite_vendue = produit_stats['quantite_vendue'] or 0
            montant_total_produit = produit_stats['montant_total_produit'] or 0

        return Response({
            "total_montant": total_montant,
            "nombre_de_ventes": total_ventes,
            "ventes": serializer.data,
            "quantite_vendue_produit": quantite_vendue,
            "montant_total_produit": montant_total_produit
        })
   



from django.db import transaction
from rest_framework.response import Response
from rest_framework import status

class VenteCacheDetailAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = VenteCache.objects.all()
    serializer_class = VenteCacheSerializer
    lookup_field = 'id'

    @transaction.atomic
    def destroy(self, request, *args, **kwargs):
        """
        Suppression d'une vente au comptant avec rollback automatique
        et restauration du stock produit.
        """
        instance = self.get_object()

        try:
            # 🔄 Restaurer le stock des produits
            for ligne in instance.lignes.all():
                produit = ligne.produit

                # Restituer la quantité
                produit.quantite += ligne.quantite

                # Restituer les codes-barres
                if ligne.codes_barres_utilises:
                    produit.codes_barres = ligne.codes_barres_utilises + produit.codes_barres

                produit.save()

            # Supprimer la vente (les lignes sont supprimées en cascade)
            self.perform_destroy(instance)

            return Response(
                {"success": "Vente supprimée avec succès et stock restauré"},
                status=status.HTTP_204_NO_CONTENT
            )

        except Exception as e:
            transaction.set_rollback(True)
            return Response(
                {"error": f"Suppression annulée : {str(e)}"},
                status=status.HTTP_400_BAD_REQUEST
            )



from docxtpl import DocxTemplate
from django.http import HttpResponse
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from .models import VenteFacilite
from django.conf import settings
from pathlib import Path
from docxtpl import DocxTemplate

class VenteFaciliteActeWordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            vente = VenteFacilite.objects.prefetch_related('lignes', 'client').get(pk=pk)
        except VenteFacilite.DoesNotExist:
            return Response({"error": "Vente non trouvée"}, status=404)

        # Charger le template ACTE.docx
        template_path = Path(settings.BASE_DIR) / "templates_word" / "ACTE.docx" # chemin vers ton modèle
        doc = DocxTemplate(template_path)

        # Construire la liste des produits
        produits_list = []
        for ligne in vente.lignes.all():
            produits_list.append({
                "designation": ligne.produit.nom,
                "quantite": ligne.quantite,
                "prix": ligne.sous_total,
            })

        # Remplir le contexte
        context = {
            "vente_id": vente.id,
            "client_nom": vente.client.nom_famille_ar,
            "client_prenom": vente.client.prenom_ar,
            "client_date_naissance":vente.client.date_naissance,
            "client_lieu_naissance":vente.client.lieu_naissance,
            "nom_pere":vente.client.nom_pere,
            "nom_mere":vente.client.nom_mere,
            "ccp":vente.client.ccp,
            "cle":vente.client.cle,
            "numero_piece_identite":vente.client.numero_piece_identite,
            "date_emission_piece":vente.client.date_emission_piece,
            "lieu_emission_piece":vente.client.lieu_emission_piece,
            "client_adresse": vente.client.adresse,
            "client_telephone": vente.client.telephone,
            "produits": produits_list,
            "total": vente.montant_total,
            "montant_verse":vente.montant_verse,
            "montant_restant":vente.montant_restant,
            "nombre_mois":vente.nombre_mois,
            "montant_mensuel":vente.montant_mensuel,
        }

        # Rendre le document
        doc.render(context)

        # Envoyer la réponse HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Acte_{vente.id}.docx"'
        doc.save(response)
        return response

class VenteFaciliteBonWordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            vente = VenteFacilite.objects.prefetch_related('lignes', 'client','utilisateur').get(pk=pk)
        except VenteFacilite.DoesNotExist:
            return Response({"error": "Vente non trouvée"}, status=404)

        # Charger le template ACTE.docx
        template_path = Path(settings.BASE_DIR) / "templates_word" / "BON.docx" # chemin vers ton modèle
        doc = DocxTemplate(template_path)

        # Construire la liste des produits
        produits_list = []
        for ligne in vente.lignes.all():
            produits_list.append({
                "designation": ligne.produit.nom,
                "prix": ligne.sous_total,
            })

        # Remplir le contexte
        context = {
            "client_nom": vente.client.nom_famille_ar,
            "client_prenom": vente.client.prenom_ar,
            "client_date_naissance":vente.client.date_naissance,
            "client_lieu_naissance":vente.client.lieu_naissance,
            "adresse":vente.client.adresse,
            "utilisateur":vente.utilisateur.username,
            "produits": produits_list,
            "total": vente.montant_total,
            "montant_verse":vente.montant_verse,
            "montant_restant":vente.montant_restant,
            "nombre_mois":vente.nombre_mois,
            "montant_mensuel":vente.montant_mensuel,
            "date_debut":vente.date_debut,
            "date_fin":vente.date_fin,
            "date_creation": vente.date_creation.strftime("%d/%m/%Y"),
        }

        # Rendre le document
        doc.render(context)

        # Envoyer la réponse HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="BON_{vente.id}.docx"'
        doc.save(response)
        return response

from pathlib import Path
from django.http import HttpResponse
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from docxtpl import DocxTemplate
from django.conf import settings
from .models import VenteCache

class VenteCashBonWordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            vente = (
                VenteCache.objects
                .select_related('client', 'utilisateur')
                .prefetch_related('lignes')
                .get(pk=pk)
            )
        except VenteCache.DoesNotExist:
            return Response({"error": "Vente non trouvée"}, status=404)

        # Load Word template
        template_path = Path(settings.BASE_DIR) / "templates_word" / "BON_CASH.docx"
        doc = DocxTemplate(template_path)

        # Build product list
        produits_list = [
            {"designation": ligne.produit.nom, "prix": f"{ligne.sous_total} DA"}
            for ligne in vente.lignes.all()
        ]

        # Fill context
        context = {
            "client_nom": vente.client.nom_famille_ar,
            "client_prenom": vente.client.prenom_ar,
            "client_date_naissance": vente.client.date_naissance.strftime("%d/%m/%Y") if vente.client.date_naissance else "",
            "client_lieu_naissance": vente.client.lieu_naissance or "",
            "adresse": vente.client.adresse or "",
            "utilisateur": vente.utilisateur.username if vente.utilisateur else "",
            "produits": produits_list,
            "total": f"{vente.montant_total} DA",
            "date_vente": vente.date_vente.strftime("%d/%m/%Y"),
        }

        # Render document
        doc.render(context)

        # Return as downloadable Word file
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="BON_{vente.id}.docx"'
        doc.save(response)
        return response


class VenteFaciliteDepotWordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        try:
            vente = VenteFacilite.objects.prefetch_related('lignes', 'client','utilisateur').get(pk=pk)
        except VenteFacilite.DoesNotExist:
            return Response({"error": "Vente non trouvée"}, status=404)

        # Charger le template ACTE.docx
        template_path = Path(settings.BASE_DIR) / "templates_word" / "DEPOT.docx" # chemin vers ton modèle
        doc = DocxTemplate(template_path)

        # Construire la liste des produits
        produits_list = []
        for ligne in vente.lignes.all():
            produits_list.append({
                "designation": ligne.produit.nom,
                "prix": ligne.sous_total,
                "magasin": ligne.produit.magasin.nom,
            })
        codes_barres_list = []

        for ligne in vente.lignes.all():
    # Ajouter tous les codes-barres de cette ligne à la liste globale
            codes_barres_list.extend(ligne.codes_barres_utilises or [])


        # Remplir le contexte
        # au lieu de passer la liste directement
        context = {
    "vente_id": vente.id,
    "client_nom": vente.client.nom_famille_ar,
    "client_prenom": vente.client.prenom_ar,
    "client_date_naissance": vente.client.date_naissance,
    "client_lieu_naissance": vente.client.lieu_naissance,
    "adresse": vente.client.adresse,
    "utilisateur": vente.utilisateur.username,
    "produits": produits_list,
    "codes_barres": "- ".join(codes_barres_list),  # <-- joint avec espace
    "date_creation": vente.date_creation.strftime("%d/%m/%Y"),
}


        # Rendre le document
        doc.render(context)

        # Envoyer la réponse HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="DEPOT_{vente.id}.docx"'
        doc.save(response)
        return response

class VenteFaciliteCCPWordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        import io
        from pathlib import Path
        from docx import Document
        from django.http import HttpResponse

        try:
            vente = (
                VenteFacilite.objects
                .select_related('client', 'utilisateur')
                .get(pk=pk)
            )
        except VenteFacilite.DoesNotExist:
            return Response({"error": "Vente non trouvée"}, status=404)

        prelevement = request.GET.get("prelevement")
        if not prelevement:
            return Response({"error": "Le paramètre 'prelevement' est requis"}, status=400)

        template_path = Path(settings.BASE_DIR) / "templates_word" / "ccp.docx"
        doc = Document(template_path)

        def remplir_ligne_cases(table, row_index, texte, is_ccp=False, cle=""):
            """
            Remplit une ligne lettre par lettre.
            Si is_ccp=True -> remplit CCP et place la clé dans les 2 dernières cellules
            """
            texte = str(texte).upper()
            cells = table.rows[row_index].cells
            total_cells = len(cells)

            # Effacer la ligne sauf la 1ère colonne (libellé)
            for i in range(1, total_cells):
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.text = ""

            # Index de la dernière colonne pour la clé
            if is_ccp:
                cle_cols = [total_cells - 2, total_cells - 1]  # 2 dernières colonnes
                data_cells = range(1, total_cells - 2)        # cellules pour CCP
            else:
                cle_cols = []
                data_cells = range(1, total_cells)

            # Remplissage CCP / texte normal
            letter_idx = 0
            for i in data_cells:
                if letter_idx < len(texte):
                    cells[i].paragraphs[0].add_run(texte[letter_idx])
                    letter_idx += 1

            # Si CCP -> remplir la clé dans les 2 dernières colonnes
            if is_ccp and cle:
                cle = cle.upper().ljust(2)[:2]  # Forcer 2 caractères max
                for idx, col in enumerate(cle_cols):
                    if idx < len(cle):
                        cells[col].paragraphs[0].add_run(cle[idx])

        # Format nombre_mois sur 2 chiffres
        nb_mois = vente.nombre_mois or 0
        echeances = f"{nb_mois:02}"

        # Données CCP
        ccp_number = str(vente.client.ccp or "")
        cle_number = str(vente.client.cle or "")

        # Remplir les deux premières tables
        for table in doc.tables[:2]:
            remplir_ligne_cases(table, 0, vente.client.nom_famille_fr or "")
            remplir_ligne_cases(table, 2, vente.client.prenom_fr or "")
            remplir_ligne_cases(table, 4, ccp_number, is_ccp=True, cle=cle_number)  # CCP + clé sur 2 colonnes
            remplir_ligne_cases(table, 8, vente.date_debut.strftime("%d"))
            remplir_ligne_cases(table, 10, vente.date_debut.strftime("%d/%m/%Y"))
            remplir_ligne_cases(table, 12, str(prelevement))
            remplir_ligne_cases(table, 14, echeances)

        # Générer le fichier Word
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="ccp_{vente.id}.docx"'
        return response
